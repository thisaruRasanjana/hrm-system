"""
generate_cvs.py — Creates 5 realistic candidate CVs as .docx files.
Run from the backend directory:
    pip install python-docx  (if not already installed)
    python seed_data/generate_cvs.py
Output folder: backend/seed_data/cvs/
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "cvs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def heading(doc, text, size=14, bold=True, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    # underline separator
    doc.add_paragraph("─" * 72)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.add_run(text).font.size = Pt(10)


def normal(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic


def job_entry(doc, title, company, dates, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{title}  —  {company}")
    r.bold = True
    r.font.size = Pt(10)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(dates)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    for b in bullets:
        bullet(doc, b)


# ─────────────────────────────────────────────────────────────────────────────
# CV 1: Kavya Sharma — Strong Python backend engineer, expect ~85–92%
# ─────────────────────────────────────────────────────────────────────────────
def cv_kavya():
    doc = Document()
    heading(doc, "Kavya Sharma", 18, color=(0x1F, 0x4E, 0x79))
    normal(doc, "kavya.sharma@email.com  |  +91 98765 43210  |  Bangalore, India  |  github.com/kavyasharma")
    doc.add_paragraph()

    section_title(doc, "Professional Summary")
    normal(doc, "Results-driven Software Engineer with 4 years of experience building scalable backend services using Python and FastAPI. Proficient in PostgreSQL, Docker, and REST API design. Experienced in Agile development environments and CI/CD pipelines. Passionate about clean, testable code and collaborative engineering.")

    section_title(doc, "Work Experience")
    job_entry(doc, "Software Engineer", "Infosys Digital (Bangalore)", "March 2022 – Present", [
        "Designed and maintained RESTful APIs using Python (FastAPI) serving 500k+ daily requests.",
        "Optimized PostgreSQL query performance, reducing average response time by 38%.",
        "Implemented CI/CD pipelines using GitHub Actions and Docker, reducing release cycle from 2 weeks to 3 days.",
        "Led migration of legacy monolithic service to microservices architecture.",
        "Mentored 2 junior engineers and conducted regular code reviews.",
    ])
    job_entry(doc, "Junior Software Engineer", "Wipro Technologies", "June 2020 – February 2022", [
        "Built internal REST APIs for HR and payroll systems using Python (Flask).",
        "Wrote unit and integration tests using pytest, achieving 87% code coverage.",
        "Worked with PostgreSQL for data modelling and report generation.",
        "Participated in Agile sprints, daily standups, and retrospectives.",
    ])

    section_title(doc, "Technical Skills")
    bullet(doc, "Languages: Python, SQL, Bash")
    bullet(doc, "Frameworks: FastAPI, Flask, SQLAlchemy")
    bullet(doc, "Databases: PostgreSQL, Redis")
    bullet(doc, "Tools: Docker, Git, GitHub Actions, Postman, Linux")
    bullet(doc, "Practices: REST API design, TDD, CI/CD, Agile/Scrum")

    section_title(doc, "Education")
    normal(doc, "B.E. Computer Science — Visvesvaraya Technological University, 2020  |  CGPA: 8.4/10")

    section_title(doc, "Certifications")
    bullet(doc, "Google Cloud Associate Cloud Engineer (2023)")
    bullet(doc, "PostgreSQL Administration — Udemy (2022)")

    path = os.path.join(OUTPUT_DIR, "01_Kavya_Sharma_CV.docx")
    doc.save(path)
    print(f"✓ Created: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# CV 2: Asel Nurlanovna — Senior Python dev, slightly overqualified, expect ~75–85%
# ─────────────────────────────────────────────────────────────────────────────
def cv_asel():
    doc = Document()
    heading(doc, "Asel Nurlanovna", 18, color=(0x1F, 0x4E, 0x79))
    normal(doc, "asel.nurlanovna@gmail.com  |  +7 701 234 5678  |  Almaty, Kazakhstan  |  linkedin.com/in/aselnurlanovna")
    doc.add_paragraph()

    section_title(doc, "Professional Summary")
    normal(doc, "Senior Software Engineer with 5+ years of experience in Python-based backend development. Specialises in API architecture, microservices, and database optimisation. Extensive experience in Agile teams across FinTech and SaaS products. Looking for a mid-level to senior role with challenging technical problems.")

    section_title(doc, "Work Experience")
    job_entry(doc, "Senior Software Engineer", "Kaspi Bank (Almaty)", "January 2021 – Present", [
        "Architected and led development of the bank's internal transaction API handling 2M+ transactions/day.",
        "Designed PostgreSQL schemas for high-throughput financial data with full ACID compliance.",
        "Introduced async Python (asyncio) patterns that improved API throughput by 60%.",
        "Conducted technical interviews and managed a team of 4 engineers.",
        "Integrated third-party payment gateways via REST APIs.",
    ])
    job_entry(doc, "Software Engineer", "Kolesa Group", "August 2019 – December 2020", [
        "Developed backend services in Python (Django REST Framework) for a marketplace platform.",
        "Collaborated with DevOps to containerize services using Docker and Kubernetes.",
        "Improved CI/CD pipeline reliability by writing automated smoke tests.",
    ])

    section_title(doc, "Technical Skills")
    bullet(doc, "Languages: Python, SQL, Go (basic)")
    bullet(doc, "Frameworks: FastAPI, Django REST Framework, Flask")
    bullet(doc, "Databases: PostgreSQL, MongoDB, Redis")
    bullet(doc, "Tools: Docker, Kubernetes, Git, Terraform, Jenkins")
    bullet(doc, "Practices: Microservices, REST & GraphQL APIs, Agile, TDD")

    section_title(doc, "Education")
    normal(doc, "B.Sc. Information Systems — Al-Farabi Kazakh National University, 2019  |  GPA: 3.9/4.0")

    section_title(doc, "Certifications")
    bullet(doc, "AWS Certified Developer – Associate (2022)")
    bullet(doc, "Professional Scrum Master I — Scrum.org (2021)")

    path = os.path.join(OUTPUT_DIR, "02_Asel_Nurlanovna_CV.docx")
    doc.save(path)
    print(f"✓ Created: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# CV 3: Daniel Fernandez — Node.js backend, limited PostgreSQL, expect ~65–75%
# ─────────────────────────────────────────────────────────────────────────────
def cv_daniel():
    doc = Document()
    heading(doc, "Daniel Fernandez", 18, color=(0x1F, 0x4E, 0x79))
    normal(doc, "daniel.fernandez@outlook.com  |  +34 612 345 678  |  Madrid, Spain  |  github.com/dani-dev")
    doc.add_paragraph()

    section_title(doc, "Professional Summary")
    normal(doc, "Software Engineer with 3 years of experience in full-stack web development with a focus on Node.js backend services and React frontends. Comfortable working across the stack, from REST API design to UI components. Experienced with Agile methodologies and Git-based team workflows.")

    section_title(doc, "Work Experience")
    job_entry(doc, "Software Engineer", "Tuenti Technologies (Madrid)", "September 2021 – Present", [
        "Built and maintained RESTful APIs using Node.js (Express) for a social communications platform.",
        "Designed and optimised MySQL database schemas; limited experience with PostgreSQL.",
        "Implemented JWT-based authentication and role-based access control.",
        "Integrated third-party APIs (Stripe payments, SendGrid email).",
        "Collaborated with frontend team on React components and API contracts.",
    ])
    job_entry(doc, "Junior Developer", "Freelance / Contract", "July 2020 – August 2021", [
        "Delivered 6 web application projects for small businesses using Node.js and MongoDB.",
        "Managed client requirements and delivered projects on schedule.",
        "Set up basic CI/CD pipelines using GitLab CI.",
    ])

    section_title(doc, "Technical Skills")
    bullet(doc, "Languages: JavaScript (ES2022), TypeScript, SQL (MySQL, some PostgreSQL)")
    bullet(doc, "Frameworks: Node.js, Express, React, Next.js")
    bullet(doc, "Databases: MySQL, MongoDB, Redis")
    bullet(doc, "Tools: Docker, Git, GitLab CI, Postman")
    bullet(doc, "Practices: REST API design, Agile, code review")

    section_title(doc, "Education")
    normal(doc, "B.Sc. Computer Engineering — Universidad Politécnica de Madrid, 2020  |  GPA: 7.8/10")

    path = os.path.join(OUTPUT_DIR, "03_Daniel_Fernandez_CV.docx")
    doc.save(path)
    print(f"✓ Created: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# CV 4: Priya Nair — Junior frontend dev, 2 years, mostly React, expect ~40–55%
# ─────────────────────────────────────────────────────────────────────────────
def cv_priya():
    doc = Document()
    heading(doc, "Priya Nair", 18, color=(0x1F, 0x4E, 0x79))
    normal(doc, "priya.nair95@gmail.com  |  +91 99887 65432  |  Kochi, India")
    doc.add_paragraph()

    section_title(doc, "Professional Summary")
    normal(doc, "Frontend Developer with 2 years of experience building responsive web applications using React and Tailwind CSS. Passionate about user experience and pixel-perfect UI implementation. Currently expanding skills into Node.js backend development. Looking for opportunities to grow into a full-stack role.")

    section_title(doc, "Work Experience")
    job_entry(doc, "Frontend Developer", "CodeCraft Solutions (Kochi)", "June 2022 – Present", [
        "Developed responsive React components for e-commerce and dashboard products.",
        "Integrated REST APIs from backend teams into frontend applications.",
        "Improved page load performance by 30% using lazy loading and code splitting.",
        "Worked closely with UI/UX designers to implement Figma designs accurately.",
        "Used Git for version control; participated in daily standups.",
    ])
    job_entry(doc, "Web Development Intern", "TechStart (Remote)", "January 2022 – May 2022", [
        "Built static websites and landing pages using HTML, CSS, and JavaScript.",
        "Assisted senior developers with React component development.",
    ])

    section_title(doc, "Technical Skills")
    bullet(doc, "Languages: JavaScript, HTML, CSS, basic Python")
    bullet(doc, "Frameworks: React, Next.js (learning), Tailwind CSS")
    bullet(doc, "Tools: Git, Figma, VS Code, npm")
    bullet(doc, "Backend: Basic Node.js/Express (self-learning), REST API consumption")

    section_title(doc, "Education")
    normal(doc, "B.Sc. Computer Applications — Mahatma Gandhi University, 2021  |  CGPA: 7.2/10")

    section_title(doc, "Projects")
    bullet(doc, "Personal Finance Tracker — React + localStorage app with charts (github.com/priyanair/fintrack)")
    bullet(doc, "E-commerce UI — Pixel-perfect implementation of a Figma design in React")

    path = os.path.join(OUTPUT_DIR, "04_Priya_Nair_CV.docx")
    doc.save(path)
    print(f"✓ Created: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# CV 5: Marcus Webb — Graphic designer pivoting to tech, expect ~15–30%
# ─────────────────────────────────────────────────────────────────────────────
def cv_marcus():
    doc = Document()
    heading(doc, "Marcus Webb", 18, color=(0x1F, 0x4E, 0x79))
    normal(doc, "marcus.webb@creativestudio.co.uk  |  +44 7911 234567  |  London, UK")
    doc.add_paragraph()

    section_title(doc, "Professional Summary")
    normal(doc, "Creative and detail-oriented Graphic Designer with 5 years of experience in brand identity, digital illustration, and UI/UX design for print and digital media. Recently completed an online Python fundamentals course and am looking to pivot into a software engineering career. Enthusiastic learner with strong visual problem-solving skills.")

    section_title(doc, "Work Experience")
    job_entry(doc, "Senior Graphic Designer", "Pixel & Ink Studio (London)", "March 2019 – Present", [
        "Led visual identity projects for 20+ brands across retail and hospitality sectors.",
        "Designed UI mockups and wireframes in Figma for client web applications.",
        "Managed junior designers and coordinated with web development teams.",
        "Produced marketing assets for digital campaigns (social media, email, web banners).",
    ])
    job_entry(doc, "Graphic Designer", "FreshCreative Agency", "July 2017 – February 2019", [
        "Created brand guidelines, logo systems, and print collateral for SME clients.",
        "Collaborated with copywriters and account managers on campaign deliverables.",
    ])

    section_title(doc, "Technical Skills")
    bullet(doc, "Design: Adobe Illustrator, Photoshop, InDesign, Figma, Canva")
    bullet(doc, "Web: Basic HTML & CSS (self-taught), Wix, WordPress")
    bullet(doc, "Programming: Python fundamentals (Udemy course, 2023), basic scripting")
    bullet(doc, "Tools: Git (basic), Slack, Trello, Google Suite")

    section_title(doc, "Education")
    normal(doc, "BA (Hons) Graphic Design — University of the Arts London, 2017  |  First Class Honours")
    normal(doc, "Python for Everybody — Coursera / University of Michigan (2023, online)")

    section_title(doc, "Note")
    normal(doc, "Actively building software engineering skills through online courses and personal projects. Completed a basic REST API project using Python and Flask. Motivated to transition into a full-time engineering role.", italic=True)

    path = os.path.join(OUTPUT_DIR, "05_Marcus_Webb_CV.docx")
    doc.save(path)
    print(f"✓ Created: {path}")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating CV files...")
    cv_kavya()
    cv_asel()
    cv_daniel()
    cv_priya()
    cv_marcus()
    print(f"\nAll done! CVs saved to: {OUTPUT_DIR}")
    print("\nUpload order (for rate limit safety — wait 5–10s between each):")
    print("  1. 01_Kavya_Sharma_CV.docx      — expect score ~85–92%")
    print("  2. 02_Asel_Nurlanovna_CV.docx   — expect score ~75–85%")
    print("  3. 03_Daniel_Fernandez_CV.docx  — expect score ~65–75%")
    print("  4. 04_Priya_Nair_CV.docx        — expect score ~40–55%")
    print("  5. 05_Marcus_Webb_CV.docx       — expect score ~15–30%")
