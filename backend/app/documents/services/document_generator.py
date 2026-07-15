"""
documents/services/document_generator.py
========================================
Service layer for generating PDF and DOCX files from templates.

Responsibilities:
- Building context dictionaries for Jinja2/DOCX placeholders.
- Merging data into HTML and DOCX templates.
- Converting HTML and DOCX output to PDF.
- Saving generated files and notifying external requesters.
"""

import io
import os
import re
import tempfile
import uuid
from datetime import datetime

from jinja2 import Template as Jinja2Template
from docx import Document
import mammoth
from xhtml2pdf import pisa
from sqlalchemy.orm import Session

from app.documents.models.request_model import DocumentRequest, RequestStatus
from app.documents.models.template_model import DocumentTemplate
from app.employees.models import Employee
from app.documents.constants import FILENAME_SUFFIX_LENGTH
from app.documents.services.email_service import send_document_to_requester
from app.core.storage_service import storage


# Ensure generated directory exists at startup



# ── Context Builders ─────────────────────────────────────────────────────────

def _build_context(employee: Employee, doc_request: DocumentRequest, override_reason: str = None) -> dict:
    """Build the template context dictionary for an internal employee request.

    Args:
        employee: The Employee ORM object.
        doc_request: The DocumentRequest ORM object.
        override_reason: Optional HR-provided text overriding the default request reason.

    Returns:
        Dictionary of variables available in templates.
    """
    reason = override_reason if override_reason else doc_request.reason
    return {
        "employee_name": f"{employee.first_name} {employee.last_name}",
        "employee_id": str(employee.id),
        "designation": getattr(employee, "designation", ""),
        "department": getattr(employee, "department", ""),
        "date": datetime.now().strftime("%B %d, %Y"),
        "reason": reason,
        "purpose": reason,  # Retained for backwards compatibility with old templates
        "document_type": doc_request.document_type,
    }


def _build_external_context(doc_request: DocumentRequest, override_reason: str = None) -> dict:
    """Build the template context dictionary for an unlinked external email request.

    Args:
        doc_request: The DocumentRequest ORM object.
        override_reason: Optional HR-provided text overriding the email subject.

    Returns:
        Dictionary of variables available in templates.
    """
    reason = override_reason if override_reason else doc_request.reason
    email_addr = doc_request.requester_email or "External Requester"
    name_part = email_addr.split("@")[0].replace(".", " ").replace("_", " ").title()
    return {
        "employee_name": name_part,
        "employee_id": "",
        "designation": "",
        "department": "",
        "date": datetime.now().strftime("%B %d, %Y"),
        "reason": reason,
        "purpose": reason,
        "document_type": doc_request.document_type,
        "requester_email": email_addr,
    }


def _replace_placeholders(text: str, context: dict) -> str:
    """Replace {{variable}} style placeholders with context values."""
    for key, value in context.items():
        text = text.replace(f"{{{{{key}}}}}", str(value))
    return text


# ── Generation Engines ───────────────────────────────────────────────────────

def _generate_from_html(template: DocumentTemplate, context: dict, preview: bool) -> tuple[str, str]:
    """Render an HTML template and convert it to PDF using xhtml2pdf.

    Returns:
        (storage_key_or_None, rendered_html_string)
    """
    try:
        jinja_template = Jinja2Template(template.content)
        rendered_html = jinja_template.render(**context)
    except Exception as e:
        raise ValueError(f"Failed to render HTML template: {e}")

    if preview:
        return None, rendered_html

    filename_base = (
        f"{context['employee_name'].replace(' ', '_')}"
        f"_{template.name.replace(' ', '_')}"
        f"_{uuid.uuid4().hex[:FILENAME_SUFFIX_LENGTH]}"
    )
    wrapper = f"""<!DOCTYPE html>
<html>
  <head>
    <style>
      @page {{ margin: 1in; }}
      body {{ font-family: 'Helvetica', 'Arial', sans-serif; font-size: 11pt; line-height: 1.4; color: #333; }}
      h1, h2, h3 {{ color: #111; margin-top: 0; }}
      p {{ margin: 0 0 12pt 0; }}
      .letter-content {{ text-align: justify; }}
    </style>
  </head>
  <body><div class="letter-content">{rendered_html}</div></body>
</html>"""

    # Try to generate a PDF from the HTML
    pdf_buf = io.BytesIO()
    try:
        pisa_status = pisa.CreatePDF(wrapper, dest=pdf_buf)
        if not pisa_status.err:
            key = f"generated_documents/{filename_base}.pdf"
            storage.upload(pdf_buf.getvalue(), key)
            return key, rendered_html
        else:
            print("Warning: xhtml2pdf encountered errors, falling back to HTML.")
    except Exception as e:
        print(f"Warning: PDF generation failed ({e}). Storing HTML instead.")

    # Fallback: store the HTML
    key = f"generated_documents/{filename_base}.html"
    storage.upload(wrapper.encode("utf-8"), key)
    return key, rendered_html


def _replace_in_paragraph(paragraph, context: dict):
    """Replace placeholders in a DOCX paragraph while preserving runs/formatting."""
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return

    replaced = _replace_placeholders(full_text, context)
    for i, run in enumerate(paragraph.runs):
        run.text = replaced if i == 0 else ""


def _generate_from_docx(template: DocumentTemplate, context: dict, preview: bool) -> tuple[str, str]:
    """Parse a DOCX file, replace placeholders, and export as PDF.

    Returns:
        (storage_key_or_None, html_preview_or_None)
    """
    if not template.file_path or not storage.file_exists(template.file_path):
        raise ValueError("DOCX template file not found.")

    # Download the DOCX template to a temp file so python-docx can open it
    template_bytes = storage.download(template.file_path)
    tmp_template = None
    tmp_docx_out = None
    try:
        suffix = os.path.splitext(template.file_path)[1] or ".docx"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tf:
            tf.write(template_bytes)
            tmp_template = tf.name

        doc = Document(tmp_template)

        for para in doc.paragraphs:
            _replace_in_paragraph(para, context)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_paragraph(para, context)
        for section in doc.sections:
            for para in section.header.paragraphs:
                _replace_in_paragraph(para, context)
            for para in section.footer.paragraphs:
                _replace_in_paragraph(para, context)

        # Save the filled DOCX to a second temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tf2:
            tmp_docx_out = tf2.name
        doc.save(tmp_docx_out)

        if preview:
            try:
                with open(tmp_docx_out, "rb") as docx_file:
                    result = mammoth.convert_to_html(docx_file)
                    html_preview = result.value
                html_preview = f"""<div style="font-family: Arial, sans-serif; line-height: 1.6; color: #333; padding: 20px;">
    {html_preview}
</div>"""
            except Exception as e:
                html_preview = f"<p>Error generating DOCX preview: {e}</p>"
            return None, html_preview

        # Convert filled DOCX to PDF
        filename_base = (
            f"{context['employee_name'].replace(' ', '_')}"
            f"_{template.name.replace(' ', '_')}"
            f"_{uuid.uuid4().hex[:FILENAME_SUFFIX_LENGTH]}"
        )
        try:
            with open(tmp_docx_out, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                raw_html = result.value

            pdf_wrapper = f"""<!DOCTYPE html>
<html>
  <head>
    <style>
      @page {{ margin: 1in; }}
      body {{ font-family: 'Helvetica', 'Arial', sans-serif; font-size: 11pt; line-height: 1.4; color: #333; }}
      p {{ margin: 0 0 12pt 0; }}
      table {{ border-collapse: collapse; width: 100%; margin-bottom: 15pt; }}
      th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
      .letter-content {{ text-align: justify; }}
    </style>
  </head>
  <body><div class="letter-content">{raw_html}</div></body>
</html>"""

            pdf_buf = io.BytesIO()
            pisa_status = pisa.CreatePDF(pdf_wrapper, dest=pdf_buf)
            if not pisa_status.err:
                key = f"generated_documents/{filename_base}.pdf"
                storage.upload(pdf_buf.getvalue(), key)
                return key, None
        except Exception as e:
            print(f"Warning: DOCX to PDF generation failed ({e}). Returning DOCX instead.")

        # Fallback: upload the filled DOCX itself
        with open(tmp_docx_out, "rb") as f:
            docx_bytes = f.read()
        key = f"generated_documents/{filename_base}.docx"
        storage.upload(docx_bytes, key)
        return key, None

    finally:
        for tmp in [tmp_template, tmp_docx_out]:
            if tmp and os.path.exists(tmp):
                os.remove(tmp)


def generate_from_custom_text(
    request_id: str, content: str, preserve_whitespace: bool = True
) -> tuple[str, str]:
    """Generate a PDF document from a raw HTML string.

    Args:
        request_id: The ID of the document request (used in filename).
        content: The raw text/HTML to place in the document body.
        preserve_whitespace: True for plain typed text; False for rich HTML.

    Returns:
        Tuple of (storage_key, raw_html_content)
    """
    white_space = "pre-wrap" if preserve_whitespace else "normal"
    html_content = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
    <style>
      @page {{ margin: 1in; }}
      body {{ font-family: 'Helvetica', 'Arial', sans-serif; font-size: 11pt; line-height: 1.4; color: #333; }}
      p {{ margin: 0 0 12pt 0; }}
      .letter-body {{ white-space: {white_space}; word-break: break-word; text-align: justify; }}
    </style>
</head>
<body>
  <div class="letter-body">{content}</div>
</body>
</html>
"""
    filename_base = f"custom_letter_{request_id}_{uuid.uuid4().hex[:FILENAME_SUFFIX_LENGTH]}"

    pdf_buf = io.BytesIO()
    try:
        pisa_status = pisa.CreatePDF(html_content, dest=pdf_buf)
        if not pisa_status.err:
            key = f"generated_documents/{filename_base}.pdf"
            storage.upload(pdf_buf.getvalue(), key)
            return key, html_content
        else:
            print("Warning: xhtml2pdf encountered errors.")
    except Exception as e:
        print(f"Warning: PDF generation failed ({e})")

    # Fallback: store HTML
    key = f"generated_documents/{filename_base}.html"
    storage.upload(html_content.encode("utf-8"), key)
    return key, html_content


# ── Main Orchestrator ────────────────────────────────────────────────────────

def _save_generated_document(db: Session, doc_request: DocumentRequest, saved_path: str) -> None:
    """Commit the final document path to the database and mark request COMPLETED."""
    final_path = saved_path.replace("\\", "/") if saved_path else None
    doc_request.status = RequestStatus.COMPLETED
    doc_request.generated_document_path = final_path
    db.commit()
    db.refresh(doc_request)

    # ── Notify requesting employee (if internal) ──────────────────────
    if doc_request.employee_id:
        try:
            from app.notifications.service import notify_employee
            notify_employee(
                db, doc_request.employee_id,
                f"Your {doc_request.document_type} document is ready to download",
                category="document", type="success", link="/dashboard/documents/request",
                entity_type="document_request", entity_id=str(doc_request.id),
            )
            db.commit()
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"[Document Generator] Notification failed: {e}")


def notify_external_requester(doc_request: DocumentRequest, final_key: str) -> None:
    """Email the generated document back to an external requester.

    Downloads the document from storage (works for both local and S3) and
    attaches it to an outbound email.
    """
    if getattr(doc_request, "source", "INTERNAL") != "EXTERNAL":
        return
    if not doc_request.requester_email or not final_key:
        return
    tmp_path = None
    try:
        # abs_path() returns a local path for local backend,
        # or a temp file path for S3 backend (caller must clean up).
        tmp_path = storage.abs_path(final_key)
        send_document_to_requester(
            to_email=doc_request.requester_email,
            document_path=tmp_path,
            document_type=doc_request.document_type,
        )
    except Exception as e:
        print(f"[Document Generator] Warning: Could not email document to requester: {e}")
    finally:
        # Clean up temp file created by S3 backend.abs_path()
        # Local backend returns a real path — do NOT delete it.
        from app.core.config import STORAGE_BACKEND
        if STORAGE_BACKEND == "s3" and tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


def generate_document_from_request(
    db: Session,
    request_id: uuid.UUID,
    template_id: int,
    preview: bool = False,
    override_reason: str = None
) -> tuple[DocumentRequest, str]:
    """Coordinate the generation of a document based on a template and request data.

    Args:
        db: Active database session.
        request_id: UUID of the request driving the generation.
        template_id: Integer ID of the template to merge.
        preview: If True, do not save to DB; just return the HTML preview.
        override_reason: Optional HR override text.

    Returns:
        Tuple of (updated DocumentRequest, generated HTML content for preview).

    Raises:
        ValueError: If the request, employee, or template are not found.
    """
    doc_request = db.query(DocumentRequest).filter(DocumentRequest.id == request_id).first()
    if not doc_request:
        raise ValueError("Document request not found")

    employee = None
    if doc_request.employee_id:
        employee = db.query(Employee).filter(Employee.id == doc_request.employee_id).first()
        if not employee:
            raise ValueError("Linked employee not found in the system")

    template = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not template:
        raise ValueError("Template not found")

    if doc_request.employee_id and employee:
        context = _build_context(employee, doc_request, override_reason)
    else:
        context = _build_external_context(doc_request, override_reason)

    template_type = (template.template_type or "").upper()

    if template_type == "HTML":
        saved_path, html_content = _generate_from_html(template, context, preview)
    elif template_type == "DOCX":
        saved_path, html_content = _generate_from_docx(template, context, preview)
    elif template_type == "PDF":
        raise ValueError(
            "PDF templates do not support variable filling. "
            "Please use an HTML or DOCX template for auto-generation."
        )
    else:
        raise ValueError(f"Unsupported template type: '{template.template_type}'")

    if preview:
        return doc_request, html_content

    _save_generated_document(db, doc_request, saved_path)
    notify_external_requester(doc_request, doc_request.generated_document_path)

    return doc_request, html_content
