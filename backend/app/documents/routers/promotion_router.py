"""
documents/routers/promotion_router.py
=======================================
Router for generating and sending Promotion Letters.

This is a standalone flow (separate from DocumentRequest) that allows HR
managers to generate a promotion letter immediately when an employee's
designation changes on the edit page.

Placeholders supported in templates:
  {{employee_name}}      – Full name of the promoted employee
  {{employee_id}}        – System employee ID (e.g. EMP-0042)
  {{old_designation}}    – Previous designation name
  {{new_designation}}    – New designation name being assigned
  {{department}}         – Employee's current department name
  {{effective_date}}     – Effective date of the promotion (HR-supplied)
  {{date}}               – Today's date (auto-set)
  {{joined_date}}        – Employee's join date
"""

import os
import uuid
from datetime import datetime, date
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, status, BackgroundTasks
from fastapi.responses import JSONResponse
from jinja2 import Template as Jinja2Template
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.database.database import get_db
from app.core.deps import require_permission
from app.employees.models import Employee, EmployeeDesignationHistory
from app.notifications.models import Notification
from app.designations.models import Designation
from app.documents.models.template_model import DocumentTemplate
from app.documents.services import document_generator
from app.documents.services.email_service import send_document_to_requester

router = APIRouter(
    prefix="/promotion-letter",
    tags=["Promotion Letters"],
)


# ── Request / Response Schemas ────────────────────────────────────────────────

class PromotionPreviewRequest(BaseModel):
    """Input for generating a promotion letter preview."""
    new_designation_id: Optional[int] = None
    effective_date: Optional[str] = None  # e.g. "2025-08-01"


class PromotionSendRequest(BaseModel):
    """Input for sending the final (possibly edited) promotion letter."""
    html_content: str                     # The edited HTML from the modal
    new_designation_id: Optional[int] = None
    effective_date: Optional[str] = None  # e.g. "2025-08-01"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _find_promotion_template(db: Session) -> DocumentTemplate:
    """Look up the Promotion Letter template by name or category (case-insensitive).

    Returns:
        The matching DocumentTemplate ORM object.

    Raises:
        HTTPException 404 with detail "TEMPLATE_MISSING" if no template found.
    """
    template = (
        db.query(DocumentTemplate)
        .filter(
            DocumentTemplate.name.ilike("%promotion%") |
            DocumentTemplate.category.ilike("%promotion%")
        )
        .first()
    )
    if not template:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="TEMPLATE_MISSING",
        )
    return template


def _build_promotion_context(
    employee: Employee,
    new_designation_name: str,
    effective_date: Optional[str],
    current_user,
) -> dict:
    """Build the template placeholder context for a promotion letter.

    The old_designation is taken from the employee's current designation string
    (the value before the HR manager saves the change).

    Args:
        employee:             The Employee ORM object.
        new_designation_name: The new designation the employee is being promoted to.
        effective_date:       HR-supplied date string for when the promotion takes effect.
        current_user:         The authenticated user making the request.

    Returns:
        Dictionary of placeholder keys → values.
    """
    dept_name = ""
    if employee.department_rel:
        dept_name = employee.department_rel.name
    elif hasattr(employee, "department") and employee.department:
        dept_name = str(employee.department)

    # Resolve the previous designation from designation history (most recent entry)
    old_designation = employee.designation or "Previous Designation"
    time_in_prev = ""
    if employee.designation_history:
        sorted_history = sorted(
            employee.designation_history,
            key=lambda h: h.start_date,
            reverse=True
        )
        if sorted_history:
            most_recent = sorted_history[0]
            old_designation = most_recent.designation_name
            if most_recent.start_date:
                delta = datetime.now().date() - most_recent.start_date
                years = delta.days // 365
                months = (delta.days % 365) // 30
                if years > 0:
                    time_in_prev = f"{years} year{'s' if years > 1 else ''}"
                    if months > 0:
                        time_in_prev += f" and {months} month{'s' if months > 1 else ''}"
                elif months > 0:
                    time_in_prev = f"{months} month{'s' if months > 1 else ''}"
                else:
                    time_in_prev = f"{delta.days} day{'s' if delta.days != 1 else ''}"
    
    signatory_name = ""
    signatory_designation = ""
    if current_user and current_user.employee:
        signatory_name = f"{current_user.employee.first_name} {current_user.employee.last_name}"
        signatory_designation = current_user.employee.designation or ""

    return {
        "Promotion Date": datetime.now().strftime("%B %d, %Y"),
        "Employee Name": f"{employee.first_name} {employee.last_name}",
        "Employee ID": employee.employee_id or str(employee.id),
        "Department": dept_name,
        "New Designation": new_designation_name,
        "Previous Designation": old_designation,
        "Time Period in Previous Designation": time_in_prev,
        "Effective Date": effective_date or datetime.now().strftime("%B %d, %Y"),
        "Authorized Signatory Name": signatory_name,
        "Designation": signatory_designation,
        
        # Legacy mappings just in case template used them
        "employee_name": f"{employee.first_name} {employee.last_name}",
    }


def _render_template(template: DocumentTemplate, context: dict) -> str:
    """Render a template's content with the given context (HTML templates only).

    For DOCX templates, this calls the existing docx-to-html rendering pipeline.

    Returns:
        Rendered HTML string.

    Raises:
        ValueError: For unsupported template types or rendering errors.
    """
    template_type = (template.template_type or "").upper()

    if template_type == "HTML":
        if not template.content:
            raise ValueError("HTML template has no content stored.")
        html = template.content
        for key, value in context.items():
            # Support exact requested placeholder format
            html = html.replace(f"[{key}]", str(value))
        
        try:
            jinja_tpl = Jinja2Template(html)
            return jinja_tpl.render(**context)
        except Exception as exc:
            raise ValueError(f"Failed to render HTML template: {exc}")

    elif template_type == "DOCX":
        if not template.file_path or not os.path.exists(template.file_path):
            raise ValueError("DOCX template file not found on disk.")
        # Build a temporary DOCX with filled placeholders, then convert to HTML for preview
        import mammoth
        from docx import Document as DocxDocument

        def _replace_in_paragraph(paragraph, ctx):
            full_text = "".join(run.text for run in paragraph.runs)
            if "{{" not in full_text and "[" not in full_text:
                return
            for key, value in ctx.items():
                full_text = full_text.replace(f"[{key}]", str(value))
                full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
            for i, run in enumerate(paragraph.runs):
                run.text = full_text if i == 0 else ""

        doc = DocxDocument(template.file_path)
        for para in doc.paragraphs:
            _replace_in_paragraph(para, context)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_paragraph(para, context)

        tmp_path = os.path.join("uploads/generated_documents", f"promo_preview_{uuid.uuid4().hex[:6]}.docx")
        doc.save(tmp_path)
        try:
            with open(tmp_path, "rb") as f:
                result = mammoth.convert_to_html(f)
                html = result.value
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

        return f'<div style="font-family: Arial, sans-serif; line-height: 1.6; color: #333; padding: 8px;">{html}</div>'

    elif template_type == "PDF":
        raise ValueError(
            "PDF templates cannot be used for promotion letters. "
            "Please use an HTML or DOCX template."
        )
    else:
        raise ValueError(f"Unsupported template type: '{template.template_type}'")


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("/{employee_id}/preview")
def preview_promotion_letter(
    employee_id: int,
    data: PromotionPreviewRequest,
    db: Session = Depends(get_db),
    current_user=Depends(require_permission("employee:update")),
):
    """Generate an HTML preview of the promotion letter for a given employee.

    Returns the rendered HTML so the frontend can display it in a preview modal
    where the HR manager can make inline edits before sending.

    Raises:
        404 TEMPLATE_MISSING – No promotion letter template exists in the system.
        404 Employee not found.
        400 Rendering error (e.g. corrupted template).
    """
    employee = (
        db.query(Employee)
        .filter(Employee.id == employee_id, Employee.is_deleted.is_(False))
        .first()
    )
    if not employee:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Employee not found.")

    new_designation_name = "New Position"
    if data.new_designation_id:
        desig = db.query(Designation).filter(Designation.id == data.new_designation_id).first()
        if desig:
            new_designation_name = desig.name

    template = _find_promotion_template(db)

    context = _build_promotion_context(employee, new_designation_name, data.effective_date, current_user)

    try:
        html_preview = _render_template(template, context)
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc))

    return {
        "preview_html":   html_preview,
        "template_name":  template.name,
        "template_type":  template.template_type,
        "employee_name":  context.get("Employee Name") or context.get("employee_name"),
        "employee_email": employee.email,
        "new_designation_name": new_designation_name,
    }


@router.post("/{employee_id}/send")
def send_promotion_letter(
    employee_id: int,
    data: PromotionSendRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user=Depends(require_permission("employee:update")),
):
    """Generate a PDF from the edited HTML and email it to the promoted employee.

    The HR manager provides the final (possibly hand-edited) HTML from the
    frontend preview. This is converted to a PDF and sent to the employee's
    registered email address.

    Returns:
        JSON with a success message and the generated document path.
    """
    employee = (
        db.query(Employee)
        .filter(Employee.id == employee_id, Employee.is_deleted.is_(False))
        .first()
    )
    if not employee:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Employee not found.")

    if not employee.email:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Employee has no email address on record.",
        )

    # Generate a unique request ID for file naming
    doc_id = uuid.uuid4().hex[:10]
    try:
        final_path, _ = document_generator.generate_from_custom_text(
            request_id=f"promotion_{employee_id}_{doc_id}",
            content=data.html_content,
            preserve_whitespace=False,   # HTML already has its own line breaks
        )
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate PDF: {exc}",
        )

    # Send the PDF to the employee's email in the background so the UI responds immediately
    def _email_task():
        try:
            abs_path = os.path.join(os.getcwd(), final_path.replace("/", os.sep))
            send_document_to_requester(
                to_email=employee.email,
                document_path=abs_path,
                document_type="Promotion Letter",
            )
        except Exception as exc:
            import logging
            logging.getLogger(__name__).error(
                f"[Promotion Letter] Failed to email letter to {employee.email}: {exc}"
            )

    # Perform the promotion if requested
    promoted = False
    if data.new_designation_id and data.new_designation_id != employee.designation_id:
        new_designation = db.query(Designation).filter_by(id=data.new_designation_id).first()
        if new_designation:
            eff_date = date.today()
            if data.effective_date:
                try:
                    eff_date = datetime.strptime(data.effective_date, "%Y-%m-%d").date()
                except ValueError:
                    pass

            # Close previous history
            if employee.designation_history:
                active_history = max(employee.designation_history, key=lambda h: (h.start_date or date.min))
                if not active_history.end_date:
                    active_history.end_date = eff_date

            # Update employee record
            employee.designation_id = data.new_designation_id
            employee.designation = new_designation.name
            
            # Start new history
            new_history = EmployeeDesignationHistory(
                employee_id=employee.id,
                designation_id=new_designation.id,
                designation_name=new_designation.name,
                start_date=eff_date,
                promotion_letter_sent=True
            )
            db.add(new_history)
            
            # Create a notification for the employee if they have an account
            if employee.user_id:
                notif = Notification(
                    user_id=employee.user_id,
                    message=f"Congratulations! You have been promoted to {new_designation.name}, effective {eff_date.strftime('%b %d, %Y')}.",
                    type="success",
                    category="system"
                )
                db.add(notif)
            
            db.commit()
            promoted = True

    if not promoted:
        # Just mark the current designation history as having the promotion letter sent
        if employee.designation_history:
            active_history = max(employee.designation_history, key=lambda h: (h.start_date or date.min))
            active_history.promotion_letter_sent = True
            db.commit()

    background_tasks.add_task(_email_task)

    return {
        "message":       "Promotion letter sent successfully.",
        "document_path": final_path,
        "sent_to":       employee.email,
    }
